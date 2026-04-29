#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_worksheet.py
生成 50 道小数加减法计算题，输出为 docx 文件（A4 单页）
适合小学 4 年级（10 岁）学生练习
"""

import random
import datetime
import sys

try:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    sys.exit(1)


# ──────────────────────────── XML 辅助 ────────────────────────────

def _remove_table_borders(table):
    """去掉表格所有可见边框"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_row_height(row, mm: float):
    """精确锁定行高（单位：毫米）"""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(mm * 56.7)))   # 1 mm ≈ 56.7 twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def _fmt(n: float) -> str:
    """把浮点数格式化为最简小数字符串，至少保留 1 位小数"""
    s = f"{n:.3f}".rstrip('0')
    if s.endswith('.'):
        s += '0'
    return s


def _fmt_num(n) -> str:
    """格式化数字：整数显示为整数，小数使用 _fmt"""
    if isinstance(n, int):
        return str(n)
    if float(n) == int(float(n)):
        return str(int(float(n)))
    return _fmt(float(n))


# ──────────────────────────── 题目生成 ────────────────────────────

def _rand_decimal(max_int: int = 15) -> float:
    """随机生成一个 1~2 位小数（适合 4 年级）"""
    dps = random.choices([1, 2], weights=[55, 45])[0]
    int_part = random.randint(0, max_int)
    if dps == 1:
        frac = random.randint(1, 9)
        return round(int_part + frac / 10, 1)
    else:
        frac = random.randint(11, 99)
        while frac % 10 == 0:          # 避免末位为 0 退化成 1 位小数
            frac = random.randint(11, 99)
        return round(int_part + frac / 100, 2)


def generate_problems(n: int = 50):
    """生成 n 道小数加减法题目（尽量不重复）"""
    seen: set = set()
    problems = []
    attempts = 0
    while len(problems) < n and attempts < n * 20:
        attempts += 1
        op = random.choice(['+', '-'])
        a = _rand_decimal()
        b = _rand_decimal()
        if op == '-':
            if a < b:
                a, b = b, a
            if abs(a - b) < 0.05:      # 避免结果接近 0
                continue
        key = (a, op, b)
        if key in seen:
            continue
        seen.add(key)
        problems.append((a, op, b))
    return problems


def generate_mul_div_problems(n: int = 50):
    """生成 n 道小数 ×/÷ 10/100/1000 题目"""
    FACTORS = [10, 100, 1000]
    seen: set = set()
    problems = []
    attempts = 0

    while len(problems) < n and attempts < n * 40:
        attempts += 1
        factor = random.choice(FACTORS)
        op = random.choice(['×', '÷'])

        if op == '×':
            # a × factor：a 为 1~2 位小数，结果 ≤ 9999
            dp = random.choice([1, 2])
            max_val = int(min(99.99, 9999 / factor) * 10 ** dp)
            if max_val < 1:
                continue
            val = random.randint(1, max_val)
            if dp == 2 and val % 10 == 0:   # 避免末位 0 退化成 1 位
                continue
            a = round(val / 10 ** dp, dp)
            b = factor
        else:
            # a ÷ factor：保证结果小数位 ≤ 3，且结果 ≥ 0.001
            if factor == 1000:
                # 整数 ÷ 1000，结果有 3 位小数
                a = float(random.randint(1, 9999))
            elif factor == 100:
                # 整数或 1 位小数 ÷ 100
                if random.random() < 0.6:
                    a = float(random.randint(1, 9999))
                else:
                    val = random.randint(11, 999)
                    if val % 10 == 0:
                        continue
                    a = round(val / 10, 1)
            else:  # factor == 10
                # 整数、1 位或 2 位小数 ÷ 10
                dp_a = random.choice([0, 1, 2])
                if dp_a == 0:
                    a = float(random.randint(1, 999))
                elif dp_a == 1:
                    a = round(random.randint(1, 99) / 10, 1)
                else:
                    val = random.randint(11, 99)
                    if val % 10 == 0:
                        continue
                    a = round(val / 100, 2)
            b = factor
            if a / b < 0.001:
                continue

        key = (a, op, b)
        if key in seen:
            continue
        seen.add(key)
        problems.append((a, op, b))

    return problems


# ──────────────────────────── 文档生成 ────────────────────────────

def _add_section_title(doc, title: str):
    """添加分节标题"""
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(3)
    r = tp.add_run(title)
    r.bold = True
    r.font.size = Pt(16)


def _add_info_row(doc):
    """添加姓名/日期/得分行"""
    ip = doc.add_paragraph()
    ip.paragraph_format.space_before = Pt(0)
    ip.paragraph_format.space_after  = Pt(5)
    today = datetime.date.today().strftime('%Y年%m月%d日')
    ir = ip.add_run(
        f"姓名：_______________   日期：{today}   "
        f"用时：_______分_______秒   得分：_______分"
    )
    ir.font.size = Pt(10)


def _add_problems_table(doc, problems: list, start_num: int = 1):
    """在 doc 末尾追加一张 2 列 × 25 行的题目表格"""
    COLS     = 2
    ROWS     = 25
    ROW_H_MM = 9.8
    COL_W    = Mm(88)

    tbl = doc.add_table(rows=ROWS, cols=COLS)
    _remove_table_borders(tbl)

    for r_idx in range(ROWS):
        row = tbl.rows[r_idx]
        _set_row_height(row, ROW_H_MM)

        for c_idx in range(COLS):
            idx  = r_idx * COLS + c_idx
            cell = row.cells[c_idx]
            cell.width = COL_W

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

            if idx < len(problems):
                a, op, b = problems[idx]

                r_num = p.add_run(f"({start_num + idx:>2})")
                r_num.bold = True
                r_num.font.size = Pt(11)

                expr = f"  {_fmt_num(a)} {op} {_fmt_num(b)} = "
                r_expr = p.add_run(expr)
                r_expr.font.size = Pt(11.5)

                r_ans = p.add_run("___________")
                r_ans.font.size = Pt(11)


def build_docx(problems: list, output_path: str, mul_div_problems: list = None):
    doc = Document()

    # ── A4 页面，窄边距，最大利用空间 ──────────────────────────────
    sec = doc.sections[0]
    sec.page_height    = Mm(297)
    sec.page_width     = Mm(210)
    sec.left_margin    = Mm(15)
    sec.right_margin   = Mm(15)
    sec.top_margin     = Mm(12)
    sec.bottom_margin  = Mm(10)

    # ── 第一节：加减法 ───────────────────────────────────────────
    _add_section_title(doc, "小数加减法计算练习（50题）")
    _add_info_row(doc)
    _add_problems_table(doc, problems, start_num=1)

    # ── 第二节：乘除法 ───────────────────────────────────────────
    if mul_div_problems:
        # 分页
        last_p = doc.add_paragraph()
        last_p.paragraph_format.space_before = Pt(0)
        last_p.paragraph_format.space_after  = Pt(0)
        from docx.oxml import OxmlElement as _OE
        run = last_p.add_run()
        br = _OE('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)

        _add_section_title(doc, "小数乘除法计算练习（50题）")
        _add_info_row(doc)
        _add_problems_table(doc, mul_div_problems, start_num=1)

    doc.save(output_path)


# ──────────────────────────── 主程序 ────────────────────────────

def main():
    add_sub = generate_problems(25)
    mul_div = generate_mul_div_problems(25)
    # 混合后打乱顺序，避免前半全是加减、后半全是乘除
    import random as _r
    problems = add_sub + mul_div
    _r.shuffle(problems)

    today       = datetime.date.today().strftime('%Y%m%d')
    output_path = f"d:/workspace/edu/小数计算练习_{today}.docx"

    build_docx(problems, output_path)

    print(f"✓ 文件已生成：{output_path}")
    print(f"  共 {len(problems)} 道题：{len(add_sub)} 道加减法 + {len(mul_div)} 道乘除法（×/÷ 10/100/1000）")
    print("  用 Word / WPS 打开，选 A4 纸、实际大小打印即可。")


if __name__ == "__main__":
    main()
